import os
import pandas as pd
<<<<<<< HEAD
import re
import time
from tqdm import tqdm
from llama_index.core import SimpleDirectoryReader, Document, Settings
from llama_index.core.node_parser import SimpleNodeParser
from llama_index.vector_stores.weaviate import WeaviateVectorStore
from llama_index.core import StorageContext
from llama_index.readers.file import PDFReader
from llama_index.core.embeddings import BaseEmbedding
from sentence_transformers import SentenceTransformer
from typing import List, Any
import weaviate
from config import (
    DATA_DIR,
    WEAVIATE_HOST,
    WEAVIATE_PORT,
    EXCEL_METADATA_FIELDS,
    CHUNK_SIZE,
    CHUNK_CLASS_PREFIX
)

def extract_incident_from_filename(filename):
    """Extract incident number from PDF filename. Returns None if not found."""
    pattern = r'^(\d{8})'  # 8 digits at the start
    match = re.match(pattern, filename)
    if match:
        return match.group(1)
    return None

def parse_excel_metadata(row, columns):
    """Parse Excel row into metadata dictionary"""
    metadata = {"document_type": "excel"}
    for field in EXCEL_METADATA_FIELDS:
        if field in columns:
            value = row[field]
            if pd.isna(value):
                metadata[field.lower().replace(" ", "_")] = ""
            else:
                metadata[field.lower().replace(" ", "_")] = str(value)
    return metadata

def ingest_excel(filepath):
    """Ingest Excel file and return list of documents"""
    try:
        df = pd.read_excel(filepath, sheet_name=None)
        docs = []
        
        for sheet_name, sheet_df in df.items():
            for _, row in sheet_df.iterrows():
                if row.isna().all():
                    continue
                    
                text_parts = []
                for col in sheet_df.columns:
                    if not pd.isna(row[col]):
                        text_parts.append(f"{col}: {row[col]}")
                
                if text_parts:
                    text = "; ".join(text_parts)
                    metadata = parse_excel_metadata(row, sheet_df.columns)
                    metadata["source"] = os.path.basename(filepath)
                    metadata["sheet_name"] = sheet_name
                    docs.append(Document(text=text, metadata=metadata))
        
        return docs
    except Exception as e:
        print(f"Error processing Excel file {filepath}: {e}")
        return []

def ingest_pdf(filepath):
    """Simplified PDF ingestion using LlamaIndex PDFReader - filename incident only"""
    filename = os.path.basename(filepath)
    
    # Extract incident number from filename ONLY
    incident_from_filename = extract_incident_from_filename(filename)
    
    # Skip PDFs that don't have incident numbers in filename
    if not incident_from_filename:
        print(f"Skipping PDF {filename} - no incident number found in filename")
        return []
    
    print(f"Processing PDF {filename} with incident number: {incident_from_filename}")
    
    try:
        # Use LlamaIndex PDFReader - handles OCR automatically
        pdf_reader = PDFReader()
        documents = pdf_reader.load_data(file=filepath)
        
        # Combine all pages into one document
        combined_text = "\n\n".join([doc.text for doc in documents])
        
        # Enhanced metadata for PDFs - using ONLY filename incident
        metadata = {
            "source": filename,
            "document_type": "pdf",
            "file_path": filepath,
            "incident_reference": incident_from_filename,  # Only from filename
            "filename_incident": incident_from_filename,   # Only from filename
            "has_filename_incident": "true",
            "extraction_method": "llamaindex_reader",
            "page_count": len(documents)
        }
        
        return [Document(text=combined_text, metadata=metadata)]
        
    except Exception as e:
        print(f"Error processing PDF {filepath}: {e}")
        return []

def ingest_docx(filepath):
    """Ingest DOCX file"""
    from docx import Document as DocxDocument
    docx = DocxDocument(filepath)
    text = "\n".join([para.text for para in docx.paragraphs])
    metadata = {
        "source": os.path.basename(filepath),
        "document_type": "docx"
    }
    return [Document(text=text, metadata=metadata)]

def ingest_directory(data_dir):
    """Ingest all supported files from directory"""
    docs = []
    pdf_processed = 0
    pdf_skipped = 0
    skipped_files = []
    processed_files = []
    
    for fname in tqdm(os.listdir(data_dir)):
        fpath = os.path.join(data_dir, fname)
        ext = fname.lower().split(".")[-1]
        
        if ext == "xlsx":
            docs.extend(ingest_excel(fpath))
            processed_files.append(f"✓ {fname} (Excel)")
        elif ext == "pdf":
            pdf_docs = ingest_pdf(fpath)
            if pdf_docs:
                docs.extend(pdf_docs)
                incident_num = extract_incident_from_filename(fname)
                processed_files.append(f"✓ {fname} (Incident: {incident_num})")
                pdf_processed += 1
            else:
                skipped_files.append(f"✗ {fname}")
                pdf_skipped += 1
        elif ext == "docx":
            docs.extend(ingest_docx(fpath))
            processed_files.append(f"✓ {fname} (DOCX)")
        elif ext in ["txt", "md"]:
            docs.extend(SimpleDirectoryReader(input_files=[fpath]).load_data())
            processed_files.append(f"✓ {fname} (Text)")
    
    # Print summary
    print(f"\n=== PDF Processing Summary ===")
    print(f"Processed PDFs with incident numbers: {pdf_processed}")
    for file in processed_files:
        if "(Incident:" in file:
            print(f"  {file}")
    print(f"Skipped PDFs without proper naming: {pdf_skipped}")
    for file in skipped_files:
        print(f"  {file}")
    print("==============================\n")
    
    return docs

def clear_weaviate_class(client, class_name):
    """Clear existing Weaviate class"""
    try:
        existing_classes = [schema['class'] for schema in client.schema.get()['classes']]
        if class_name in existing_classes:
            print(f"Class '{class_name}' exists. Deleting...")
            client.schema.delete_class(class_name)
            time.sleep(2)
            print(f"Class '{class_name}' deleted.")
        else:
            print(f"Class '{class_name}' does not exist. Proceeding with ingestion.")
    except Exception as e:
        print(f"Error during class cleanup: {e}")

class LocalSentenceTransformerEmbedding(BaseEmbedding):
    """Local embedding using SentenceTransformer"""
    _model: Any = None
    
    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        print("Loading SentenceTransformer model...")
        object.__setattr__(self, '_model', SentenceTransformer("sentence-transformers/all-MiniLM-L6-v2"))
        print("SentenceTransformer model loaded successfully")
        
    def _get_query_embedding(self, query: str) -> List[float]:
        return self._model.encode(query).tolist()
        
    def _get_text_embedding(self, text: str) -> List[float]:
        return self._model.encode(text).tolist()
        
    def _get_text_embeddings(self, texts: List[str]) -> List[List[float]]:
        return [self._model.encode(text).tolist() for text in texts]
        
    async def _aget_query_embedding(self, query: str) -> List[float]:
        return self._get_query_embedding(query)
        
    async def _aget_text_embedding(self, text: str) -> List[float]:
        return self._get_text_embedding(text)
        
    async def _aget_text_embeddings(self, texts: List[str]) -> List[List[float]]:
        return self._get_text_embeddings(texts)

def get_class_name():
    """Get current class name based on date"""
    import datetime
    timestamp = datetime.datetime.now().strftime("%Y%m%d")
    return f"{CHUNK_CLASS_PREFIX}_{timestamp}"

def save_class_name(class_name):
    """Save class name to a file for main.py to read"""
    with open("current_class.txt", "w") as f:
        f.write(class_name)
    print(f"Class name saved to current_class.txt: {class_name}")

def main():
    """Main ingestion function"""
    # Setup embedding model
    Settings.embed_model = LocalSentenceTransformerEmbedding()
    print("Using local SentenceTransformer embedding")
    
    # Ingest documents
    docs = ingest_directory(DATA_DIR)
    print(f"Ingested {len(docs)} documents. Indexing to Weaviate...")

    # Connect to Weaviate
    client = weaviate.Client(url=f"http://{WEAVIATE_HOST}:{WEAVIATE_PORT}")
    
    # Use date-based class name
    class_name = get_class_name()
    print(f"Using class name: {class_name}")
    
    # Clear existing class and create vector store
    clear_weaviate_class(client, class_name)
    vector_store = WeaviateVectorStore(weaviate_client=client, index_name=class_name)

    # Process documents into nodes
    parser = SimpleNodeParser(chunk_size=CHUNK_SIZE)
    nodes = parser.get_nodes_from_documents(docs)
    
    print(f"Processing {len(nodes)} nodes...")
    
    # Generate embeddings and store in Weaviate
    for i, node in enumerate(nodes):
        if node.text and node.text.strip():
            # Generate embedding
            node.embedding = Settings.embed_model._get_text_embedding(node.text)
            
            # Clean metadata
            clean_metadata = {}
            for key, value in node.metadata.items():
                clean_metadata[key] = str(value) if value is not None else ""
            
            # Store in Weaviate
            try:
                client.data_object.create(
                    data_object={
                        "text": node.text,
                        **clean_metadata
                    },
                    class_name=class_name,
                    vector=node.embedding
                )
                if (i + 1) % 5 == 0:
                    print(f"Inserted {i + 1}/{len(nodes)} nodes")
            except Exception as e:
                print(f"Error inserting node {i}: {e}")

    print(f"Ingestion complete! Stored {len(nodes)} nodes in class: {class_name}")
    
    # Save class name for main.py
    save_class_name(class_name)
    
    return class_name
=======
from sentence_transformers import SentenceTransformer
import weaviate
from config import EMBEDDING_MODEL, WEAVIATE_CLASS_NAME, WEAVIATE_URL

import pdfplumber
import docx

def parse_to_chunks(file_path):
    """Parse the file at file_path into text chunks depending on type (Excel, PDF, Word), including column headers in chunk text for tabular data."""
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()

    chunks = []

    if ext in [".xlsx", ".xls"]:
        with pd.ExcelFile(file_path) as xls:
            for sheet_name in xls.sheet_names:
                df = pd.read_excel(xls, sheet_name=sheet_name)
                headers = df.columns.tolist()
                for idx, row in df.iterrows():
                    row_data = {header: row[header] for header in headers}
                    chunk_text = ", ".join([f"{k}: {v}" for k, v in row_data.items() if pd.notnull(v)])
                    if chunk_text.strip():
                        chunks.append({
                            "text": chunk_text,
                            "sheet": sheet_name,
                            "row": idx,
                            "source_file": os.path.basename(file_path)
                        })
    elif ext == ".pdf":
        with pdfplumber.open(file_path) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text()
                if text and text.strip():
                    paragraphs = [p for p in text.split('\n') if p.strip()]
                    for j, para in enumerate(paragraphs):
                        if para.strip():
                            chunks.append({
                                "text": para.strip(),
                                "sheet": f"page_{i+1}",
                                "row": j,
                                "source_file": os.path.basename(file_path)
                            })
    elif ext == ".docx":
        doc = docx.Document(file_path)
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                chunks.append({
                    "text": text,
                    "sheet": "docx",
                    "row": i,
                    "source_file": os.path.basename(file_path)
                })
    else:
        raise ValueError(f"Unsupported file type: {ext}")

    if not chunks:
        raise ValueError(f"No data found in {file_path}.")

    return chunks

def parse_all_data_folder():
    """Detects all supported files in data/ and parses them to chunks."""
    data_dir = os.path.abspath("data")
    if not os.path.exists(data_dir):
        raise FileNotFoundError(f"Data folder not found: {data_dir}")
    chunks = []
    for fname in os.listdir(data_dir):
        if fname.lower().endswith(('.xlsx', '.xls', '.pdf', '.docx')):
            file_path = os.path.join(data_dir, fname)
            chunks.extend(parse_to_chunks(file_path))
    if not chunks:
        raise ValueError("No supported files found or no data parsed in the data folder.")
    return chunks

def ingest_to_weaviate(chunks):
    """Ingest chunks into Weaviate v4."""
    client = weaviate.Client("http://localhost:8080")
    existing_classes = [cls["class"] for cls in client.schema.get()["classes"]]
    if WEAVIATE_CLASS_NAME not in existing_classes:
        class_obj = {
            "class": WEAVIATE_CLASS_NAME,
            "properties": [
                {"name": "text", "dataType": ["text"]},
                {"name": "sheet", "dataType": ["text"]},
                {"name": "row", "dataType": ["int"]}
            ]
        }
        client.schema.create_class(class_obj)
    model = SentenceTransformer(EMBEDDING_MODEL)
    for chunk in chunks:
        embedding = model.encode([chunk["text"]])[0].tolist()
        resp = client.data_object.create(
            data_object={
                "text": chunk["text"],
                "sheet": chunk["sheet"],
                "row": int(chunk["row"])
            },
            class_name=WEAVIATE_CLASS_NAME,
            vector=embedding
        )
    return len(chunks)

def main():
    chunks = parse_all_data_folder()
    count = ingest_to_weaviate(chunks)
    print(f"✅ Ingested {count} chunks from all supported files in 'data' into Weaviate.")
>>>>>>> 376396da3a0af583750614f36a751c757cc6dddc

if __name__ == "__main__":
    main()